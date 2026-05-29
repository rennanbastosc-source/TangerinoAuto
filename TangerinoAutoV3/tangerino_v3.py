import subprocess, sys, threading, queue, traceback, json, time, ctypes, os
from datetime import datetime
from pathlib import Path

# Registra AppUserModelID antes de criar qualquer janela para que o Windows
# mostre o ícone correto na barra de tarefas (não o ícone do python.exe)
try:
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("Tangerino.Auto.V2")
except Exception:
    pass

# Detecta se está rodando como executável PyInstaller
_FROZEN = getattr(sys, "frozen", False)
_BASE   = Path(sys._MEIPASS) if _FROZEN else Path(__file__).parent

# Quando congelado, direciona o Chromium e o cache para AppData do usuário
# (único lugar gravável num exe portátil)
if _FROZEN:
    _app_data = Path(os.environ.get("APPDATA", Path.home())) / "TangerinoAuto"
    _app_data.mkdir(parents=True, exist_ok=True)
    os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(_app_data / "browsers")

# ── abre a janela imediatamente ───────────────────────────────────────────────
if not _FROZEN:
    for _pkg in ["customtkinter", "tkcalendar"]:
        try:
            __import__(_pkg)
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", _pkg],
                                  stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

import customtkinter as ctk
from tkcalendar import DateEntry

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("green")  # base neutra — cores primárias definidas via LARANJA

LARANJA        = "#E8630A"
LARANJA_HOVER  = "#C4520A"
LARANJA_SUBTIT = "#FFDBB5"

# ── configurações ─────────────────────────────────────────────────────────────

STATIC_AUTH = ""                    # capturado automaticamente no login
USER_TYPE   = "ADMINISTRATOR"
_runtime    = {"static_auth": ""}   # preenchido pelo Playwright durante o login
REPORT_API  = "https://report.tangerino.com.br/async-reports"
WS_BASE     = "wss://report.tangerino.com.br/websocket"
LOGIN_URL   = "https://app.tangerino.com.br/Tangerino/pages/LoginPage"

def _decoded_jwt(token):
    """Decodifica o payload do JWT (sem verificar assinatura) e retorna o dict."""
    import base64, json as _json
    try:
        payload_b64 = token.split(".")[1]
        payload_b64 += "=" * (-len(payload_b64) % 4)  # padding
        return _json.loads(base64.urlsafe_b64decode(payload_b64))
    except Exception:
        return {}

def _ws_dest(token):
    """Retorna o destino WebSocket correto para o usuário logado."""
    claims = _decoded_jwt(token)
    # JWT do Tangerino usa 'userId' como ID numérico do usuário
    user_id   = claims.get("userId") or 0
    user_type = claims.get("userType") or USER_TYPE
    return f"/app/session/report/time-sheet/{user_type}_{user_id}"

def _user_name(token):
    """Retorna o nome do usuário extraído do JWT."""
    claims = _decoded_jwt(token)
    # JWT do Tangerino usa 'userName' para o nome completo
    return claims.get("userName") or claims.get("name") or "ADMINISTRATOR"

FOLHA_URL   = "https://app.tangerino.com.br/Tangerino/pages/folha-ponto?funcionalidade=24&wicket:pageMapName=wicket-0"

PASTA_DESTINO = Path.home() / "Documents" / "Conferencia Diaria V3"
PASTA_DESTINO.mkdir(exist_ok=True)
PASTA_LOGS    = PASTA_DESTINO / "logs"
PASTA_LOGS.mkdir(exist_ok=True)

EXCLUIDOS        = {"RENANN BASTOS CAVALCANTE"}
HORARIO_OVERRIDE = {"ALEXANDRA MARQUES DE LIMA": "07:00-17:00"}
TOLERANCIA       = 10
DIAS_SEMANA      = {0: "Segunda", 1: "Ter", 2: "Quarta", 3: "Quinta", 4: "Sexta", 5: "S", 6: "Dom"}
COR_HEADER       = (31, 73, 125)
COR_ALT          = (217, 226, 243)

_mods  = {}
_token = {"value": None}
_creds = {"usuario": None, "senha": None}

CACHE_FILE      = (_app_data if _FROZEN else Path(__file__).parent) / ".cache.json"
TOKEN_TTL       = 23 * 3600   # token válido por 23h (margem antes das 24h)
LOCAIS_TTL      = 7 * 86400   # locais raramente mudam — cache de 7 dias

def _ler_cache():
    try:
        return json.loads(CACHE_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _salvar_cache(dados: dict):
    try:
        atual = _ler_cache()
        atual.update(dados)
        CACHE_FILE.write_text(json.dumps(atual), encoding="utf-8")
    except Exception:
        pass

def _cache_valido(chave, ttl):
    c = _ler_cache()
    ts = c.get(f"{chave}_ts", 0)
    return c.get(chave) if (time.time() - ts) < ttl else None

def _invalidar_token():
    """Remove apenas o token do cache, preservando locais."""
    c = _ler_cache()
    c.pop("token", None)
    c.pop("token_ts", None)
    try:
        CACHE_FILE.write_text(json.dumps(c), encoding="utf-8")
    except Exception:
        pass

def _invalidar_locais():
    """Remove locais do cache, forçando nova busca."""
    c = _ler_cache()
    c.pop("locais", None)
    c.pop("locais_ts", None)
    try:
        CACHE_FILE.write_text(json.dumps(c), encoding="utf-8")
    except Exception:
        pass

def _renovar_token(log):
    """Faz novo login e atualiza o cache. Retorna o novo token."""
    log("[Auth] Token inválido — renovando automaticamente...")
    token = obter_token(log)
    if token:
        _salvar_cache({"token": token, "token_ts": time.time()})
        _token["value"] = token
        log("[Auth] Token renovado com sucesso.")
    else:
        log("[Auth] Falha ao renovar token.")
    return token


# ── inicialização em background ───────────────────────────────────────────────

def _inicializar(log, on_locais_prontos, on_erro):
    try:
        # 1. pacotes
        pkgs = {"requests": "requests", "websocket": "websocket-client",
                "playwright": "playwright", "pdfplumber": "pdfplumber", "docx": "python-docx"}
        for mod, pkg in pkgs.items():
            try:
                _mods[mod] = __import__(mod)
            except ImportError:
                log(f"[Init] Instalando {pkg}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", pkg],
                                      stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                _mods[mod] = __import__(mod)

        from playwright.sync_api import sync_playwright
        _mods["sync_playwright"] = sync_playwright

        if _FROZEN:
            # No exe portátil usa node(.exe) + cli.js embutidos pelo PyInstaller
            import platform as _plat
            _node_bin = "node.exe" if _plat.system() == "Windows" else "node"
            _node = _BASE / "playwright" / "driver" / _node_bin
            _cli  = _BASE / "playwright" / "driver" / "package" / "cli.js"
            if _node.exists() and _cli.exists():
                log("[Init] Baixando Chromium (primeira execução)...")
                result = subprocess.run(
                    [str(_node), str(_cli), "install", "chromium"],
                    capture_output=True, check=False,
                    env={**os.environ},
                )
                if result.returncode != 0:
                    log(f"[Init] Aviso Chromium: {result.stderr.decode(errors='ignore')[:300]}")
            else:
                log(f"[Init] Driver não localizado em {_node}")
        else:
            subprocess.run([sys.executable, "-m", "playwright", "install", "chromium", "--quiet"],
                           capture_output=True)

        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        _mods.update({"Pt": Pt, "RGBColor": RGBColor, "Cm": Cm,
                      "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
                      "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
                      "WD_ORIENT": WD_ORIENT, "qn": qn, "OxmlElement": OxmlElement})

        # 2. token — usa cache se ainda válido
        token = _cache_valido("token", TOKEN_TTL)
        cached_auth = _ler_cache().get("static_auth", "")
        if token and cached_auth:
            log("[Login] Token em cache reutilizado (sem login necessário).")
            _runtime["static_auth"] = cached_auth
        else:
            log("[Login] Abrindo navegador headless...")
            token = obter_token(log)
            if not token:
                on_erro("Falha no login.")
                return
            _salvar_cache({"token": token, "token_ts": time.time(),
                           "static_auth": _runtime["static_auth"]})
        _token["value"] = token

        # 3. locais — usa cache se ainda válido
        LOCAIS_VER = 2  # incrementar sempre que a lógica de filtragem mudar
        locais = _cache_valido("locais", LOCAIS_TTL)
        if locais and _ler_cache().get("locais_ver", 1) < LOCAIS_VER:
            locais = None  # cache desatualizado, rebusca
            _invalidar_locais()
        if locais:
            log(f"[Locais] {len(locais)} local(is) carregado(s) do cache.")
        else:
            log("[Locais] Buscando locais de trabalho...")
            locais = buscar_locais(token, log)
            if not locais:
                on_erro("Nenhum local encontrado. Veja o log acima.")
                return
            _salvar_cache({"locais": locais, "locais_ts": time.time(), "locais_ver": LOCAIS_VER})

        on_locais_prontos(locais)

    except Exception:
        log(f"[Init] ERRO:\n{traceback.format_exc()}")
        on_erro("Erro na inicialização. Veja o log.")


# ── token (idêntico à v1) ─────────────────────────────────────────────────────

def obter_token(log):
    import json as _json

    token_holder  = [None]
    locais_holder = [None]

    with _mods["sync_playwright"]() as p:
        browser = p.chromium.launch(headless=True, args=[
            "--no-sandbox", "--disable-gpu",
            "--disable-dev-shm-usage", "--disable-extensions",
        ])
        page = browser.new_page()

        # bloqueia recursos que não precisamos (imagens, fontes, mídia)
        def bloquear(route):
            if route.request.resource_type in ("image", "font", "media", "stylesheet"):
                route.abort()
            else:
                route.continue_()
        page.route("**/*", bloquear)

        def on_request(request):
            t = request.headers.get("tng-web-token")
            if t and not token_holder[0]:
                token_holder[0] = t
            # captura o STATIC_AUTH (authorization) dos requests ao domínio de relatórios
            if "report.tangerino.com.br" in request.url and not _runtime["static_auth"]:
                auth = request.headers.get("authorization", "")
                if auth:
                    _runtime["static_auth"] = auth
                    log(f"[Login] STATIC_AUTH capturado automaticamente.")

        def on_response(response):
            url = response.url
            # loga toda URL chamada após o login para descobrir endpoint de locais
            if "tangerino" in url and any(x in url for x in (
                "work", "place", "local", "obra", "empresa", "employer", "location"
            )):
                log(f"[DEBUG] {response.status} {url}")
                try:
                    body = response.text()
                    log(f"[DEBUG] body: {body[:300]}")
                except Exception:
                    pass

            # tenta capturar locais de qualquer resposta JSON com lista de objetos
            if locais_holder[0]:
                return
            ct = response.headers.get("content-type", "")
            if "json" not in ct:
                return
            try:
                dados = response.json()
                if isinstance(dados, dict):
                    for chave in ("content", "data", "workPlaces", "workplaces", "result", "items"):
                        if isinstance(dados.get(chave), list) and dados[chave]:
                            dados = dados[chave]
                            break
                if isinstance(dados, list) and dados and isinstance(dados[0], dict) and "id" in dados[0]:
                    campos_nome = ("name", "description", "fantasyName", "razaoSocial")
                    campo = next((c for c in campos_nome if dados[0].get(c)), None)
                    if campo:
                        locais = {str(item[campo]): item["id"] for item in dados if item.get("id")}
                        if locais:
                            log(f"[Locais] Capturado via intercepção: {len(locais)} local(is) — URL: {url}")
                            locais_holder[0] = locais
            except Exception:
                pass

        page.on("request",  on_request)
        page.on("response", on_response)

        page.goto(LOGIN_URL, wait_until="domcontentloaded")
        try:
            page.locator("text=Empregador").first.click()
            page.wait_for_timeout(800)
        except Exception:
            pass
        page.fill('input[name="login"]', _creds["usuario"])
        page.fill('input[name="password"]', _creds["senha"])
        page.click('input[name="btnLogin"]')
        try:
            page.wait_for_load_state("networkidle", timeout=15000)
        except Exception:
            pass
        page.goto(FOLHA_URL, wait_until="domcontentloaded")

        # aguarda token + locais (máx 20s)
        for _ in range(40):
            if token_holder[0] and locais_holder[0]:
                break
            page.wait_for_timeout(500)

        # se não capturou locais ainda, aguarda mais um pouco (dropdown pode demorar)
        if token_holder[0] and not locais_holder[0]:
            log("[Locais] Aguardando carregamento do dropdown...")
            for _ in range(20):
                if locais_holder[0]:
                    break
                page.wait_for_timeout(500)

        browser.close()

    if token_holder[0]:
        log("[Login] Token capturado com sucesso.")
    else:
        log("[Login] Token não encontrado.")

    # guarda locais capturados para uso posterior
    _token["locais"] = locais_holder[0] or {}
    return token_holder[0]


# ── busca de locais com log completo ──────────────────────────────────────────

def buscar_locais(token, log):
    """Busca todos os locais paginando /filter/workplaces. Renova token se necessário."""
    base    = "https://report.tangerino.com.br/filter/workplaces"
    locais  = {}
    inativas = 0
    pagina  = 1

    while True:
        url  = f"{base}?page={pagina}&size=100"
        resp = _mods["requests"].get(url, headers=_headers_api(token), timeout=15)

        if resp.status_code in (401, 403):
            token = _renovar_token(log)
            if not token:
                break
            resp = _mods["requests"].get(url, headers=_headers_api(token), timeout=15)

        if resp.status_code != 200:
            log(f"[Locais] Status inesperado {resp.status_code}.")
            break

        lista = resp.json().get("list", [])
        if not lista:
            break
        for item in lista:
            if not item.get("active", True):
                inativas += 1
                continue
            id_  = item.get("id")
            nome = item.get("name") or str(id_)
            if id_:
                locais[nome] = id_
        if len(lista) < 100:
            break
        pagina += 1

    log(f"[Locais] {len(locais)} ativo(s) carregado(s), {inativas} inativo(s) ignorado(s).")
    return locais


# ── relatório + WebSocket ─────────────────────────────────────────────────────

def _headers_api(token):
    auth = _runtime["static_auth"] or STATIC_AUTH
    return {
        "tng-web-token": token, "authorization": auth,
        "Content-Type": "application/json",
        "Origin":  "https://report-web.tangerino.com.br",
        "Referer": "https://report-web.tangerino.com.br/",
    }

def solicitar_relatorio(token, data_ini, data_fim, workplace_id, log):
    payload = {
        "authorization": STATIC_AUTH,
        "filter": {
            "employee": None, "employeeId": None,
            "company": None,  "companyId": None,
            "workPlace": None, "workPlaceId": workplace_id,
            "manager": None,  "managerId": None,
            "jobRole": None,  "jobRoleId": None,
            "costCenter": None,
            "startDate": data_ini, "endDate": data_fim,
            "format": "PDF", "statusEmployee": "ADMITIDOS",
            "showHourBalance": True, "showOnlyEmployeesWithPunch": True,
            "showDsr": False, "showHorasInterjornada": False,
            "showHorasIntrajornada": False,
        },
        "type": "TIME_SHEET",
        "user": {"name": _user_name(token), "type": USER_TYPE},
    }
    for tentativa in range(2):
        resp = _mods["requests"].post(REPORT_API, json=payload,
                                      headers=_headers_api(token), timeout=30)
        log(f"[API] Status: {resp.status_code}")
        if resp.status_code == 200:
            return True, token
        log(f"[API] Resposta: {resp.text[:300]}")
        if resp.status_code in (401, 403) and tentativa == 0:
            token = _renovar_token(log)
            if not token:
                return False, token
            payload  # token atualizado na próxima iteração
            continue
        return False, token
    return False, token


def ws_url():
    import random, string
    sid = "".join(random.choices(string.ascii_lowercase + string.digits, k=8))
    return f"{WS_BASE}/{random.randint(0, 999)}/{sid}/websocket"


def aguardar_pdf(token, log):
    import json as _json, ssl
    resultado = {"url": None}
    pronto    = threading.Event()

    def on_message(ws, msg):
        if msg == "o":
            ws.send(_json.dumps([
                f"CONNECT\nAuthorization:Bearer {token}\naccept-version:1.0,1.1,1.2\nheart-beat:10000,10000\n\n\x00"
            ]))
        elif msg.startswith("a"):
            for frame in _json.loads(msg[1:]):
                if frame.startswith("CONNECTED"):
                    dest = _ws_dest(token)
                    log(f"[WS] Inscrito em: {dest}")
                    ws.send(_json.dumps([f"SUBSCRIBE\nid:sub-1\ndestination:{dest}\n\n\x00"]))
                    log("[WS] Aguardando PDF...")
                elif "MESSAGE" in frame and "time-sheet" in frame:
                    body = frame[frame.find("\n\n") + 2:].rstrip("\x00")
                    try:
                        url = _json.loads(body).get("content", {}).get("fileUrl")
                        if url:
                            resultado["url"] = url
                            pronto.set()
                            ws.close()
                    except Exception:
                        pass

    def on_error(ws, err):
        log(f"[WS] Erro: {err}")
        pronto.set()

    ws_app = _mods["websocket"].WebSocketApp(ws_url(), on_message=on_message, on_error=on_error)
    threading.Thread(target=ws_app.run_forever,
                     kwargs={"sslopt": {"cert_reqs": ssl.CERT_NONE}}, daemon=True).start()
    pronto.wait(timeout=90)
    return resultado["url"]


def _nome_seguro(texto, limite=60):
    """Remove caracteres inválidos para nomes de pasta no Windows."""
    import re
    texto = re.sub(r'[\\/:*?"<>|]', '-', texto)
    return texto[:limite].strip()

def criar_pasta_entrega(nome_local, label_data):
    """Cria subpasta: Conferencia Diaria V2 / DD-MM-YYYY - OBRA - HHhMM"""
    horario = datetime.now().strftime("%Hh%M")
    data_pasta = label_data.replace("/", "-")
    nome_pasta = f"{data_pasta} - {_nome_seguro(nome_local)} - {horario}"
    pasta = PASTA_DESTINO / nome_pasta
    pasta.mkdir(parents=True, exist_ok=True)
    return pasta

def baixar_pdf(url, data_display, pasta, log):
    resp    = _mods["requests"].get(url, timeout=60)
    nome    = f"folha_ponto_{data_display.replace('/', '-')}.pdf"
    destino = pasta / nome
    with open(destino, "wb") as f:
        f.write(resp.content)
    log(f"[PDF] Salvo: {nome}")
    return destino


# ── parser PDF + gerador DOCX (idêntico à v1) ────────────────────────────────

def t2m(s):
    neg = s.strip().startswith('-')
    h, m = map(int, s.lstrip('-').strip().split(':'))
    return -(h * 60 + m) if neg else h * 60 + m

def m2s(mins):
    h, m = divmod(abs(mins), 60)
    return f"{'-' if mins < 0 else ''}{h}:{m:02d}"

def horario_previsto(texto, data_str):
    import re
    try:
        prefix = DIAS_SEMANA[datetime.strptime(data_str, "%d/%m/%Y").weekday()]
    except Exception:
        return []
    for linha in texto.splitlines():
        if re.match(rf"\s*{prefix}", linha, re.IGNORECASE):
            pares = re.findall(r'(\d{2}:\d{2})\s+\S{1,3}\s+(\d{2}:\d{2})', linha)
            if pares:
                return [f"{ini}-{fim}" for ini, fim in pares]
    return []

def _linha_do_dia(texto, data_str):
    """Retorna a linha do PDF que corresponde à data DD/MM informada."""
    import re
    try:
        dia_mes = datetime.strptime(data_str, "%d/%m/%Y").strftime("%d/%m")
    except Exception:
        return ""
    for linha in texto.splitlines():
        if re.match(rf"\s*{re.escape(dia_mes)}\b", linha):
            return linha
    return ""

def pontos_do_dia(texto, data_str=None):
    import re
    padrao = r'\d{2}/\d{2}\s+\S+\s+((?:\d{2}:\d{2}\s+)*)\|\s*(\d{2}:\d{2})\s+(\d{2}:\d{2})\s+(-?\d+:\d+)'
    # se data_str informada, filtra pela linha específica do dia
    alvo = _linha_do_dia(texto, data_str) if data_str else texto
    m = re.search(padrao, alvo) if alvo else None
    if not m:
        return "", 0, 0, 0
    pontos = re.findall(r'\d{2}:\d{2}', m.group(1))
    pares  = " / ".join(f"{pontos[i]}-{pontos[i+1]}" for i in range(0, len(pontos)-1, 2))
    return pares or "-", t2m(m.group(2)), t2m(m.group(3)), t2m(m.group(4))

def extrair_ocorrencias(caminho_pdf, data_str):
    import re
    ocorrencias = []
    with _mods["pdfplumber"].open(caminho_pdf) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ""
            m = re.search(r"Nome:\s+(.+?)\s+CPF:", texto)
            if not m:
                continue
            nome = m.group(1).strip()
            if nome in EXCLUIDOS:
                continue
            funcao = ""
            m = re.search(r"Fun.{1,3}o:\s+(.+?)\s+Centro de Custo:", texto)
            if m:
                funcao = m.group(1).strip()
            previsto_str = HORARIO_OVERRIDE.get(nome) or \
                           " / ".join(horario_previsto(texto, data_str)) or "-"
            linha_dia = _linha_do_dia(texto, data_str)
            if "FALTA NAO JUSTIFICA" in (linha_dia or texto):
                ocorrencias.append({"nome": nome, "funcao": funcao,
                                    "previsto": previsto_str, "registrado": "—",
                                    "ocorrencia": "Falta"})
                continue
            registrado, trab, prev, deficit = pontos_do_dia(texto, data_str)
            if deficit == 0 or abs(deficit) <= TOLERANCIA:
                continue
            ocorrencia = "Ausência parcial" if abs(deficit) >= 60 else f"Atraso {m2s(deficit)}"
            ocorrencias.append({"nome": nome, "funcao": funcao,
                                "previsto": previsto_str, "registrado": registrado,
                                "ocorrencia": ocorrencia})
    return ocorrencias

def set_cell_bg(cell, rgb):
    shd = _mods["OxmlElement"]("w:shd")
    shd.set(_mods["qn"]("w:val"), "clear")
    shd.set(_mods["qn"]("w:color"), "auto")
    shd.set(_mods["qn"]("w:fill"), "{:02X}{:02X}{:02X}".format(*rgb))
    cell._tc.get_or_add_tcPr().append(shd)

def cell_write(cell, texto, negrito=False, cor_txt=None, pt=9, alinhar="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = (_mods["WD_ALIGN_PARAGRAPH"].CENTER
                   if alinhar == "center" else _mods["WD_ALIGN_PARAGRAPH"].LEFT)
    run = p.add_run(texto)
    run.bold = negrito
    run.font.size = _mods["Pt"](pt)
    if cor_txt:
        run.font.color.rgb = _mods["RGBColor"](*cor_txt)

def _gerar_pagina_docx(doc, data_str, nome_local, ocorrencias):
    """Adiciona uma página (cabeçalho + tabela) ao documento para uma data específica."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    Pt = _mods["Pt"]; RGBColor = _mods["RGBColor"]; Cm = _mods["Cm"]
    WA = _mods["WD_ALIGN_PARAGRAPH"]; WT = _mods["WD_TABLE_ALIGNMENT"]

    t = doc.add_paragraph(); t.alignment = WA.CENTER
    r = t.add_run("REGISTRO DE OCORRÊNCIAS DE PONTO")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(*COR_HEADER)

    s = doc.add_paragraph(); s.alignment = WA.CENTER
    s.add_run(f"Local: {nome_local}").font.size = Pt(10)

    dp = doc.add_paragraph(); dp.alignment = WA.CENTER
    rp = dp.add_run(f"Data de referência: {data_str}")
    rp.font.size = Pt(10); rp.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()

    colunas  = ["Nº", "Colaborador", "Função", "Horário Previsto", "Horário Registrado", "Ocorrência", "Justificativa"]
    larguras = [Cm(0.9), Cm(5.5), Cm(4), Cm(4.5), Cm(4), Cm(3.5), Cm(3.3)]
    centrar  = {0, 4, 5}

    tabela = doc.add_table(rows=1, cols=len(colunas))
    tabela.alignment = WT.CENTER; tabela.style = "Table Grid"

    for i, (col, larg) in enumerate(zip(colunas, larguras)):
        c = tabela.rows[0].cells[i]; c.width = larg
        set_cell_bg(c, COR_HEADER)
        cell_write(c, col, negrito=True, cor_txt=(255,255,255), pt=9,
                   alinhar="center" if i in centrar else "left")

    for idx, oc in enumerate(ocorrencias, 1):
        cor = COR_ALT if idx % 2 == 0 else (255, 255, 255)
        row = tabela.add_row()
        vals = [str(idx), oc["nome"], oc["funcao"], oc["previsto"], oc["registrado"], oc["ocorrencia"], ""]
        for i, (val, larg) in enumerate(zip(vals, larguras)):
            c = row.cells[i]; c.width = larg
            set_cell_bg(c, cor)
            cell_write(c, val, pt=9, alinhar="center" if i in centrar else "left")

    doc.add_paragraph()
    rod = doc.add_paragraph()
    rr  = rod.add_run("* Preencha o campo Justificativa e encaminhe para aprovação.")
    rr.italic = True; rr.font.size = Pt(8); rr.font.color.rgb = RGBColor(100, 100, 100)


def gerar_docx(ocs_por_data, label, nome_local, pasta, log):
    """ocs_por_data: dict {data_str: [ocorrencias]} ordenado cronologicamente."""
    from docx import Document
    Pt = _mods["Pt"]; RGBColor = _mods["RGBColor"]; Cm = _mods["Cm"]
    WO = _mods["WD_ORIENT"]

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WO.LANDSCAPE
    sec.page_width  = Cm(29.7); sec.page_height   = Cm(21)
    sec.top_margin  = Cm(1.8);  sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2);    sec.right_margin  = Cm(2)

    datas = list(ocs_por_data.items())
    for i, (data_str, ocorrencias) in enumerate(datas):
        if i > 0:
            # quebra de página entre datas
            doc.add_page_break()
        _gerar_pagina_docx(doc, data_str, nome_local, ocorrencias)

    arq = f"justificativas_{label.replace('/', '-').replace(' a ', '_a_')}.docx"
    caminho = pasta / arq
    doc.save(str(caminho))
    log(f"[DOCX] Salvo: {arq} ({len(datas)} página(s))")
    return caminho


# ── interface gráfica ─────────────────────────────────────────────────────────

class TangerinoApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Tangerino Auto v3.0")
        self.resizable(False, False)
        try:
            import platform as _plat
            if _plat.system() == "Windows":
                _ico = _BASE / "tangerino.ico"
                if _ico.exists():
                    self.iconbitmap(default=str(_ico))
            else:
                import tkinter as _tk
                _png = _BASE / "tangerino.png"
                if _png.exists():
                    _photo = _tk.PhotoImage(file=str(_png))
                    self.iconphoto(True, _photo)
                    self._icon_ref = _photo
        except Exception:
            pass
        self._log_queue  = queue.Queue()
        self._all_logs   = []   # acumula todos os logs para diagnóstico
        self._rodando    = False
        self._local_vars = {}
        self._poll_log()
        self._build_login_ui()

    def _centralizar(self, w, h):
        """Define tamanho e posiciona a janela no centro do monitor."""
        self.update_idletasks()
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        x  = (sw - w) // 2
        y  = (sh - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    # ── tela de login ─────────────────────────────────────────────────────────

    def _build_login_ui(self):
        self._centralizar(440, 460)
        self._frame_login = ctk.CTkFrame(self, fg_color="transparent")
        self._frame_login.pack(fill="both", expand=True, padx=36, pady=28)

        header = ctk.CTkFrame(self._frame_login, fg_color=LARANJA, corner_radius=10)
        header.pack(fill="x", pady=(0, 22))
        ctk.CTkLabel(header, text="Tangerino Auto",
                     font=ctk.CTkFont(size=22, weight="bold"), text_color="white").pack(pady=(14,2))
        ctk.CTkLabel(header, text="v3.0 — Gerador de Folha de Ponto",
                     font=ctk.CTkFont(size=11), text_color=LARANJA_SUBTIT).pack(pady=(0,14))

        ctk.CTkLabel(self._frame_login, text="E-mail",
                     font=ctk.CTkFont(size=12, weight="bold"), anchor="w").pack(fill="x")
        self._ent_usuario = ctk.CTkEntry(self._frame_login, height=36,
                                          placeholder_text="seu@email.com")
        self._ent_usuario.pack(fill="x", pady=(4, 12))

        ctk.CTkLabel(self._frame_login, text="Senha",
                     font=ctk.CTkFont(size=12, weight="bold"), anchor="w").pack(fill="x")
        row_senha = ctk.CTkFrame(self._frame_login, fg_color="transparent")
        row_senha.pack(fill="x", pady=(4, 6))
        row_senha.columnconfigure(0, weight=1)
        self._ent_senha = ctk.CTkEntry(row_senha, height=36, show="•",
                                        placeholder_text="••••••••")
        self._ent_senha.grid(row=0, column=0, sticky="ew", padx=(0, 8))
        self._var_show = ctk.BooleanVar(value=False)
        ctk.CTkCheckBox(row_senha, text="Mostrar", variable=self._var_show,
                        font=ctk.CTkFont(size=11), width=80,
                        command=lambda: self._ent_senha.configure(
                            show="" if self._var_show.get() else "•")
                        ).grid(row=0, column=1)

        self._var_lembrar = ctk.BooleanVar(value=False)
        ctk.CTkCheckBox(self._frame_login, text="Lembrar login neste computador",
                        variable=self._var_lembrar,
                        font=ctk.CTkFont(size=11)).pack(anchor="w", pady=(8, 16))

        self._btn_entrar = ctk.CTkButton(self._frame_login, text="Entrar", height=42,
                                          font=ctk.CTkFont(size=14, weight="bold"),
                                          fg_color=LARANJA, hover_color=LARANJA_HOVER,
                                          command=self._fazer_login)
        self._btn_entrar.pack(fill="x")

        self._lbl_status = ctk.CTkLabel(self._frame_login, text="",
                                         font=ctk.CTkFont(size=11), text_color="#FF6B6B")
        self._lbl_status.pack(pady=(8, 0))

        # pré-preenche credenciais salvas
        creds = _ler_cache().get("credenciais")
        if creds:
            self._ent_usuario.insert(0, creds.get("usuario", ""))
            self._ent_senha.insert(0, creds.get("senha", ""))
            self._var_lembrar.set(True)

        self._ent_usuario.bind("<Return>", lambda e: self._ent_senha.focus())
        self._ent_senha.bind("<Return>",   lambda e: self._fazer_login())
        self._ent_usuario.focus()

    def _fazer_login(self):
        usuario = self._ent_usuario.get().strip()
        senha   = self._ent_senha.get()
        if not usuario or not senha:
            self._lbl_status.configure(text="⚠  Preencha e-mail e senha.", text_color="#FFA500")
            return

        _creds["usuario"] = usuario
        _creds["senha"]   = senha

        if self._var_lembrar.get():
            _salvar_cache({"credenciais": {"usuario": usuario, "senha": senha}})
        else:
            c = _ler_cache()
            c.pop("credenciais", None)
            try:
                CACHE_FILE.write_text(json.dumps(c), encoding="utf-8")
            except Exception:
                pass

        self._btn_entrar.configure(state="disabled", text="Conectando...")
        self._lbl_status.configure(text="Aguarde, isso pode levar até 30 segundos...", text_color="gray")

        threading.Thread(target=_inicializar,
                         args=(self._log, self._on_login_prontos, self._on_login_erro),
                         daemon=True).start()

    def _on_login_prontos(self, locais):
        self.after(0, lambda: self._ir_para_principal(locais))

    def _on_login_erro(self, msg):
        self.after(0, lambda: self._btn_entrar.configure(state="normal", text="Entrar"))
        self.after(0, lambda: self._lbl_status.configure(text=f"✗  {msg}", text_color="#FF6B6B"))
        self.after(0, lambda: self._mostrar_log_erro())

    def _salvar_log_execucao(self):
        """Salva o log da execução atual em Conferencia Diaria V3/logs/."""
        try:
            ts   = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
            nome = f"log_{ts}.txt"
            conteudo = "\n".join(self._all_logs)
            (PASTA_LOGS / nome).write_text(conteudo, encoding="utf-8")
        except Exception:
            pass

    def _mostrar_log_erro(self, titulo="Erro de inicialização", conteudo_extra=""):
        conteudo = "\n".join(self._all_logs)
        if conteudo_extra:
            conteudo += "\n" + conteudo_extra

        # salva na pasta de logs da aplicação
        try:
            ts   = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
            nome = f"erro_{ts}.txt"
            log_file = PASTA_LOGS / nome
            log_file.write_text(conteudo, encoding="utf-8")
        except Exception:
            log_file = PASTA_LOGS / "erro.txt"

        # popup com o conteúdo
        popup = ctk.CTkToplevel(self)
        popup.title("⚠  " + titulo)
        popup.geometry("600x420")
        popup.resizable(True, True)
        popup.grab_set()

        ctk.CTkLabel(popup, text=titulo + ":",
                     font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=16, pady=(12,4))

        box = ctk.CTkTextbox(popup, font=ctk.CTkFont(family="Consolas", size=10))
        box.pack(fill="both", expand=True, padx=16, pady=(0,8))
        box.insert("end", conteudo)
        box.configure(state="disabled")

        ctk.CTkLabel(popup, text=f"Log salvo em: {log_file}",
                     font=ctk.CTkFont(size=10), text_color="gray").pack(anchor="w", padx=16)
        ctk.CTkButton(popup, text="Fechar", width=100, height=32,
                      command=popup.destroy).pack(pady=(6,12))

    def _ir_para_principal(self, locais):
        self._frame_login.destroy()
        self._centralizar(640, 640)
        self._build_ui()
        self._popular(locais)
        self._btn.configure(state="normal", text="▶  Gerar Relatório")
        self._log("[Init] Conectado com sucesso.")

    # ── tela principal ────────────────────────────────────────────────────────

    def _build_ui(self):
        header = ctk.CTkFrame(self, fg_color=LARANJA, corner_radius=10)
        header.pack(fill="x", padx=20, pady=(20,10))
        ctk.CTkLabel(header, text="Tangerino Auto",
                     font=ctk.CTkFont(size=22, weight="bold"), text_color="white").pack(pady=(12,2))
        ctk.CTkLabel(header, text="v3.0 — Gerador de Folha de Ponto",
                     font=ctk.CTkFont(size=12), text_color=LARANJA_SUBTIT).pack(pady=(0,12))

        ctk.CTkLabel(self, text="Local de Trabalho",
                     font=ctk.CTkFont(size=13, weight="bold"), anchor="w").pack(fill="x", padx=24, pady=(10,4))

        self._busca_var = ctk.StringVar()
        self._busca_var.trace_add("write", lambda *_: self._filtrar_locais())
        self._busca_entry = ctk.CTkEntry(self, textvariable=self._busca_var,
                                          placeholder_text="🔍  Pesquisar obra...",
                                          height=32, font=ctk.CTkFont(size=12))
        self._busca_entry.pack(fill="x", padx=20, pady=(0,4))

        self._locais_frame = ctk.CTkScrollableFrame(self, corner_radius=8, height=120)
        self._locais_frame.pack(fill="x", padx=20, pady=(0,4))
        self._placeholder = ctk.CTkLabel(self._locais_frame,
                                         text="Conectando ao Tangerino...",
                                         text_color="gray", font=ctk.CTkFont(size=12))
        self._placeholder.pack(anchor="w", padx=16, pady=10)

        sel = ctk.CTkFrame(self, fg_color="transparent")
        sel.pack(fill="x", padx=20, pady=(0,8))
        ctk.CTkButton(sel, text="Selecionar todos", width=140, height=28,
                      font=ctk.CTkFont(size=11), fg_color=LARANJA, hover_color=LARANJA_HOVER,
                      command=lambda: self._sel_todos(True)).pack(side="left", padx=(0,8))
        ctk.CTkButton(sel, text="Desmarcar todos", width=140, height=28,
                      font=ctk.CTkFont(size=11), fg_color="gray60", hover_color="gray50",
                      command=lambda: self._sel_todos(False)).pack(side="left", padx=(0,8))

        self._tabs = ctk.CTkTabview(self, height=80, corner_radius=8)
        self._tabs.pack(fill="x", padx=20, pady=(4, 10))
        self._tabs.add("Dia Único")
        self._tabs.add("Período")

        cal_opts = dict(
            date_pattern="dd/mm/yyyy",
            maxdate=datetime.today().date(),
            background=LARANJA, foreground="white",
            headersbackground=LARANJA, headersforeground="white",
            selectbackground=LARANJA_HOVER, selectforeground="white",
            normalbackground="white", normalforeground="black",
            weekendbackground="#F5F5F5", weekendforeground="#888888",
            othermonthbackground="#EEEEEE", othermonthforeground="#AAAAAA",
            bordercolor="#DDDDDD", font=("Arial", 11),
        )

        # aba Dia Único
        aba_dia = self._tabs.tab("Dia Único")
        aba_dia.columnconfigure(0, weight=1)
        self._cal_dia = DateEntry(aba_dia, width=22, **cal_opts)
        self._cal_dia.set_date(datetime.today().date())
        self._cal_dia.grid(row=0, column=0, padx=(0, 6), pady=10, sticky="ew")
        ctk.CTkButton(aba_dia, text="Hoje", width=70, height=32,
                      fg_color=LARANJA, hover_color=LARANJA_HOVER,
                      command=self._set_hoje).grid(row=0, column=1, pady=10)

        # aba Período
        aba_per = self._tabs.tab("Período")
        aba_per.columnconfigure(1, weight=1)
        aba_per.columnconfigure(3, weight=1)
        ctk.CTkLabel(aba_per, text="De:", font=ctk.CTkFont(size=12)).grid(
            row=0, column=0, padx=(0, 4), pady=10)
        self._cal_ini = DateEntry(aba_per, width=16, **cal_opts)
        self._cal_ini.set_date(datetime.today().date())
        self._cal_ini.grid(row=0, column=1, padx=(0, 12), pady=10, sticky="ew")
        ctk.CTkLabel(aba_per, text="Até:", font=ctk.CTkFont(size=12)).grid(
            row=0, column=2, padx=(0, 4), pady=10)
        self._cal_fim = DateEntry(aba_per, width=16, **cal_opts)
        self._cal_fim.set_date(datetime.today().date())
        self._cal_fim.grid(row=0, column=3, pady=10, sticky="ew")

        self._btn = ctk.CTkButton(self, text="Inicializando...", height=44,
                                  font=ctk.CTkFont(size=14, weight="bold"),
                                  fg_color=LARANJA, hover_color=LARANJA_HOVER,
                                  state="disabled", command=self._iniciar)
        self._btn.pack(fill="x", padx=20, pady=(4,10))

        log_header = ctk.CTkFrame(self, fg_color="transparent")
        log_header.pack(fill="x", padx=20)
        ctk.CTkButton(log_header, text="⏻  Sair", width=90, height=26,
                      font=ctk.CTkFont(size=11), fg_color="gray60", hover_color="gray50",
                      command=self._logout).pack(side="right", padx=(8,0))
        ctk.CTkButton(log_header, text="📋 Copiar Log", width=110, height=26,
                      font=ctk.CTkFont(size=11), fg_color="gray60", hover_color="gray50",
                      command=self._copiar_log).pack(side="right")

        self._log_box = ctk.CTkTextbox(self, font=ctk.CTkFont(family="Consolas", size=11),
                                       height=180, state="disabled")
        self._log_box.pack(fill="both", expand=True, padx=20, pady=(4,20))

    def _on_locais_prontos(self, locais):
        self.after(0, lambda: self._popular(locais))
        self.after(0, lambda: self._btn.configure(state="normal", text="▶  Gerar Relatório"))

    def _on_erro(self, msg):
        self._log(f"✗  {msg}")
        self.after(0, lambda: self._placeholder.configure(text=f"✗ {msg}", text_color="#FF6B6B"))

    def _popular(self, locais):
        self._placeholder.destroy()
        self._checkboxes = {}  # nome → widget CTkCheckBox
        for nome, id_ in sorted(locais.items()):
            var = ctk.BooleanVar(value=False)   # desmarcadas por padrão
            self._local_vars[nome] = (var, id_)
            cb = ctk.CTkCheckBox(self._locais_frame, text=nome, variable=var,
                                 font=ctk.CTkFont(size=11))
            cb.pack(anchor="w", padx=16, pady=3)
            self._checkboxes[nome] = cb

    def _filtrar_locais(self):
        termo = self._busca_var.get().lower().strip()
        for nome, cb in self._checkboxes.items():
            if termo in nome.lower():
                cb.pack(anchor="w", padx=16, pady=3)
            else:
                cb.pack_forget()

    def _forcar_login(self):
        """Apaga o cache de token e reinicia a inicialização."""
        if self._rodando:
            return
        _invalidar_token()
        _invalidar_locais()
        self._btn.configure(state="disabled", text="Reconectando...")
        for widget in self._locais_frame.winfo_children():
            widget.destroy()
        self._placeholder = ctk.CTkLabel(self._locais_frame,
                                         text="Reconectando...",
                                         text_color="gray", font=ctk.CTkFont(size=12))
        self._placeholder.pack(anchor="w", padx=16, pady=10)
        self._local_vars.clear()
        self._checkboxes = {}
        self._busca_var.set("")
        self._log("[Auth] Cache de token apagado. Fazendo novo login...")
        threading.Thread(target=_inicializar,
                         args=(self._log, self._on_locais_prontos, self._on_erro),
                         daemon=True).start()

    def _copiar_log(self):
        conteudo = self._log_box.get("1.0", "end").strip()
        if not conteudo:
            self._log("⚠  Log vazio, nada para copiar.")
            return
        self.clipboard_clear()
        self.clipboard_append(conteudo)
        self._log("✓  Log copiado para a área de transferência.")

    def _logout(self):
        """Limpa token e volta para a tela de login sem reiniciar o processo."""
        if self._rodando:
            self._log("⚠  Aguarde o processo atual terminar para sair.")
            return
        _invalidar_token()
        _invalidar_locais()
        _runtime["static_auth"] = _ler_cache().get("static_auth", "")
        _token["value"] = None
        # remove token e static_auth do cache (mantém credenciais salvas)
        try:
            import json as _j
            cache = _ler_cache()
            for k in ("token", "token_ts"):
                cache.pop(k, None)
            CACHE_FILE.write_text(_j.dumps(cache), encoding="utf-8")
        except Exception:
            pass
        # destrói todos os widgets atuais e reconstrói a tela de login
        for w in self.winfo_children():
            w.destroy()
        # remove referência ao log_box para o _poll_log não tentar acessá-lo
        if hasattr(self, "_log_box"):
            del self._log_box
        self._all_logs  = []
        self._log_queue = queue.Queue()
        self._rodando   = False
        self._local_vars = {}
        self._build_login_ui()

    def _sel_todos(self, v):
        for var, _ in self._local_vars.values():
            var.set(v)

    def _set_hoje(self):
        self._cal_dia.set_date(datetime.today().date())

    def _log(self, msg):
        self._all_logs.append(msg)
        self._log_queue.put(msg)

    def _poll_log(self):
        if hasattr(self, "_log_box"):
            try:
                while not self._log_queue.empty():
                    msg = self._log_queue.get_nowait()
                    self._log_box.configure(state="normal")
                    self._log_box.insert("end", msg + "\n")
                    self._log_box.see("end")
                    self._log_box.configure(state="disabled")
            except Exception:
                pass
        self.after(100, self._poll_log)

    def _iniciar(self):
        if self._rodando:
            return
        locais = {n: id_ for n, (var, id_) in self._local_vars.items() if var.get()}
        if not locais:
            self._log("⚠  Selecione ao menos um local.")
            return

        aba = self._tabs.get()
        if aba == "Dia Único":
            data_ini = data_fim = self._cal_dia.get_date()
        else:
            data_ini = self._cal_ini.get_date()
            data_fim = self._cal_fim.get_date()
            if data_fim < data_ini:
                self._log("⚠  Data final anterior à data inicial.")
                return
        self._rodando = True
        self._btn.configure(state="disabled", text="Aguarde...")
        self._log_box.configure(state="normal")
        self._log_box.delete("1.0", "end")
        self._log_box.configure(state="disabled")
        threading.Thread(
            target=self._executar,
            args=(locais, data_ini.strftime("%d/%m/%Y"), data_fim.strftime("%d/%m/%Y")),
            daemon=True
        ).start()

    def _executar(self, locais, data_ini_str, data_fim_str):
        ini_api = datetime.strptime(data_ini_str, "%d/%m/%Y").strftime("%Y-%m-%d")
        fim_api = datetime.strptime(data_fim_str, "%d/%m/%Y").strftime("%Y-%m-%d")
        label   = data_ini_str if data_ini_str == data_fim_str else f"{data_ini_str} a {data_fim_str}"
        token   = _token["value"]
        try:
            for nome_local, wid in locais.items():
                self._log(f"\n── {nome_local} ({label}) ──")
                pdf_url = [None]

                def ws_worker():
                    pdf_url[0] = aguardar_pdf(token, self._log)

                t = threading.Thread(target=ws_worker, daemon=True)
                t.start(); time.sleep(2)

                ok, token = solicitar_relatorio(token, ini_api, fim_api, wid, self._log)
                if not ok:
                    self._log("✗  Falha ao solicitar relatório.")
                    self._salvar_log_execucao()
                    self.after(0, lambda n=nome_local: self._mostrar_log_erro(
                        titulo=f"Falha ao solicitar relatório — {n}"))
                    continue
                t.join(timeout=90)
                if not pdf_url[0]:
                    self._log("✗  PDF não recebido (timeout 90s).")
                    self._salvar_log_execucao()
                    self.after(0, lambda n=nome_local: self._mostrar_log_erro(
                        titulo=f"PDF não recebido — {n}"))
                    continue

                pasta = criar_pasta_entrega(nome_local, label)
                self._log(f"[Pasta] {pasta.name}")

                self._log("[Download] Baixando PDF...")
                caminho = baixar_pdf(pdf_url[0], data_fim_str, pasta, self._log)

                self._log("[DOCX] Gerando...")
                from datetime import timedelta
                ini_dt = datetime.strptime(data_ini_str, "%d/%m/%Y")
                fim_dt = datetime.strptime(data_fim_str, "%d/%m/%Y")
                ocs_por_data = {}
                d = ini_dt
                while d <= fim_dt:
                    ds  = d.strftime("%d/%m/%Y")
                    ocs = extrair_ocorrencias(caminho, ds)
                    if ocs:
                        ocs_por_data[ds] = ocs
                    d += timedelta(days=1)
                total = sum(len(v) for v in ocs_por_data.values())
                self._log(f"[DOCX] {total} ocorrência(s) em {len(ocs_por_data)} dia(s) com pendência.")
                if ocs_por_data:
                    gerar_docx(ocs_por_data, label, nome_local, pasta, self._log)
                else:
                    self._log("[DOCX] Sem ocorrências — arquivo não gerado.")

            self._log("\n✓  Concluído!")
            self._salvar_log_execucao()
            os.startfile(str(PASTA_DESTINO))
        except Exception:
            tb = traceback.format_exc()
            self._log(f"✗  Erro:\n{tb}")
            self._salvar_log_execucao()
            self.after(0, lambda: self._mostrar_log_erro(
                titulo="Erro inesperado na emissão", conteudo_extra=tb))
        finally:
            self._rodando = False
            self.after(0, lambda: self._btn.configure(state="normal", text="▶  Gerar Relatório"))


if __name__ == "__main__":
    app = TangerinoApp()
    app.mainloop()
