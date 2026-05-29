import subprocess, sys, os, re
from pathlib import Path
from datetime import datetime

for pkg in ["pdfplumber", "python-docx"]:
    try:
        __import__(pkg.replace("-", "_").replace("python_", ""))
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PASTA_PDF  = Path.home() / "Documents" / "conferencia diária"
PASTA_DOCX = PASTA_PDF
NOME_LOCAL  = "CRECHE E ESCOLA DE EDUCAÇÃO INFANTIL, CEI PADRÃO FNDE DISTRITO DE COITE"
EXCLUIDOS   = {"RENANN BASTOS CAVALCANTE"}  # empregador, desobrigado de registrar ponto

# Horários que diferem do que está no PDF (sobrescreve a leitura automática)
HORARIO_OVERRIDE = {
    "ALEXANDRA MARQUES DE LIMA": "07:00-17:00",
}
COR_HEADER = (31, 73, 125)
COR_ALT    = (217, 226, 243)
TOLERANCIA = 10  # minutos — déficit <= 10 min não é ocorrência

DIAS_SEMANA = {0: "Segunda", 1: "Ter", 2: "Quarta", 3: "Quinta", 4: "Sexta", 5: "S", 6: "Dom"}


# ── utils de tempo ────────────────────────────────────────────────────────────

def t2m(s):
    """'07:30' ou '-3:50' → int minutos"""
    s = s.strip()
    neg = s.startswith('-')
    h, m = map(int, s.lstrip('-').split(':'))
    return -(h * 60 + m) if neg else h * 60 + m

def m2s(mins):
    """int minutos → '-3:50'"""
    neg = mins < 0
    h, m = divmod(abs(mins), 60)
    return f"{'-' if neg else ''}{h}:{m:02d}"


# ── parsing do PDF ────────────────────────────────────────────────────────────

def horario_previsto(texto, data_str):
    """
    Lê o Quadro de Horários e retorna os períodos do dia indicado.
    Ex: ['07:00-12:00', '13:00-17:00']
    """
    try:
        d = datetime.strptime(data_str, "%d/%m/%Y")
        prefix = DIAS_SEMANA[d.weekday()]
    except Exception:
        return []

    for linha in texto.splitlines():
        if re.match(rf"\s*{prefix}", linha, re.IGNORECASE):
            # "Quinta-feira 07:00 às 12:00 13:00 às 17:00 09:00"
            # "às" pode estar corrompido em qualquer coisa de 1-3 chars
            pares = re.findall(r'(\d{2}:\d{2})\s+\S{1,3}\s+(\d{2}:\d{2})', linha)
            if pares:
                return [f"{ini}-{fim}" for ini, fim in pares]
    return []


def pontos_do_dia(texto):
    """
    Linha: '28/05 quinta-feira 06:56 12:06 | 05:10 09:00 -3:50'
    Retorna: (pontos_str, trabalhadas_min, previstas_min, deficit_min)
    """
    m = re.search(
        r'\d{2}/\d{2}\s+\S+-\w+\s+((?:\d{2}:\d{2}\s+)*)\|\s*(\d{2}:\d{2})\s+(\d{2}:\d{2})\s+(-?\d+:\d+)',
        texto
    )
    if not m:
        return "", 0, 0, 0

    pontos  = re.findall(r'\d{2}:\d{2}', m.group(1))
    trab    = t2m(m.group(2))
    prev    = t2m(m.group(3))
    deficit = t2m(m.group(4))

    # Formata pares entrada-saída
    pares = " / ".join(
        f"{pontos[i]}-{pontos[i+1]}" for i in range(0, len(pontos) - 1, 2)
    )
    return pares or "-", trab, prev, deficit


def extrair_ocorrencias(caminho_pdf, data_str=None):
    ocorrencias = []

    with pdfplumber.open(caminho_pdf) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ""

            # Nome
            m = re.search(r"Nome:\s+(.+?)\s+CPF:", texto)
            if not m:
                continue
            nome = m.group(1).strip()
            if nome in EXCLUIDOS:
                continue

            # Função
            funcao = ""
            m = re.search(r"Fun.{1,3}o:\s+(.+?)\s+Centro de Custo:", texto)
            if m:
                funcao = m.group(1).strip()

            if nome in HORARIO_OVERRIDE:
                previsto_str = HORARIO_OVERRIDE[nome]
            else:
                previsto_lista = horario_previsto(texto, data_str) if data_str else []
                previsto_str   = " / ".join(previsto_lista) if previsto_lista else "-"

            # ── Falta total ──
            if "FALTA NAO JUSTIFICADA" in texto:
                ocorrencias.append({
                    "nome": nome, "funcao": funcao,
                    "previsto": previsto_str,
                    "registrado": "—",
                    "ocorrencia": "Falta",
                })
                continue

            # ── Presença com déficit ──
            registrado, trabalhadas, previstas, deficit = pontos_do_dia(texto)

            # Dentro da tolerância → sem ocorrência
            if deficit == 0 or abs(deficit) <= TOLERANCIA:
                continue

            # Classifica o déficit
            if abs(deficit) >= 60:
                # Déficit grande: não cumpriu um dos períodos do dia
                ocorrencia = "Ausência parcial"
            else:
                # Déficit pequeno: atraso ou saída ligeiramente antecipada
                ocorrencia = f"Atraso {m2s(deficit)}"

            ocorrencias.append({
                "nome": nome, "funcao": funcao,
                "previsto": previsto_str,
                "registrado": registrado,
                "ocorrencia": ocorrencia,
            })

    return ocorrencias


# ── geração do DOCX ───────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_tuple):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*rgb_tuple))
    cell._tc.get_or_add_tcPr().append(shd)


def cell_write(cell, texto, negrito=False, cor_txt=None, pt=9, alinhar="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(alinhar, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(texto)
    run.bold = negrito
    run.font.size = Pt(pt)
    if cor_txt:
        run.font.color.rgb = RGBColor(*cor_txt)


def gerar_docx(ocorrencias, data_str, caminho_saida):
    doc = Document()

    # Paisagem A4
    sec = doc.sections[0]
    sec.orientation  = WD_ORIENT.LANDSCAPE
    sec.page_width   = Cm(29.7)
    sec.page_height  = Cm(21)
    sec.top_margin   = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin  = Cm(2)
    sec.right_margin = Cm(2)

    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("REGISTRO DE OCORRÊNCIAS DE PONTO")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(*COR_HEADER)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"Local: {NOME_LOCAL}").font.size = Pt(10)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = dp.add_run(f"Data de referência: {data_str}")
    rp.font.size = Pt(10); rp.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Colunas  (total ~25.7cm com margens de 2cm cada lado)
    colunas  = ["Nº", "Colaborador", "Função", "Horário Previsto", "Horário Registrado", "Ocorrência", "Justificativa"]
    larguras = [Cm(0.9), Cm(5.5), Cm(4), Cm(4.5), Cm(4), Cm(3.5), Cm(3.3)]
    centrar  = {0, 4, 5}  # índices que ficam centralizados

    tabela = doc.add_table(rows=1, cols=len(colunas))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Table Grid"

    # Cabeçalho
    hdr = tabela.rows[0]
    for i, (col, larg) in enumerate(zip(colunas, larguras)):
        c = hdr.cells[i]
        c.width = larg
        set_cell_bg(c, COR_HEADER)
        cell_write(c, col, negrito=True, cor_txt=(255, 255, 255), pt=9,
                   alinhar="center" if i in centrar else "left")

    # Dados
    for idx, oc in enumerate(ocorrencias, 1):
        cor = COR_ALT if idx % 2 == 0 else (255, 255, 255)
        row = tabela.add_row()
        valores = [
            str(idx),
            oc["nome"],
            oc["funcao"],
            oc["previsto"],
            oc["registrado"],
            oc["ocorrencia"],
            "",
        ]
        for i, (val, larg) in enumerate(zip(valores, larguras)):
            c = row.cells[i]
            c.width = larg
            set_cell_bg(c, cor)
            cell_write(c, val, pt=9, alinhar="center" if i in centrar else "left")

    doc.add_paragraph()

    rod = doc.add_paragraph()
    rr = rod.add_run("* Preencha o campo Justificativa e encaminhe para aprovação.")
    rr.italic = True; rr.font.size = Pt(8); rr.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(str(caminho_saida))
    print(f"[DOCX] Salvo em: {caminho_saida}")


# ── main ──────────────────────────────────────────────────────────────────────

def pdf_mais_recente():
    pdfs = sorted(PASTA_PDF.glob("folha_ponto_*.pdf"), key=os.path.getmtime, reverse=True)
    if not pdfs:
        raise FileNotFoundError(f"Nenhum PDF encontrado em {PASTA_PDF}")
    return pdfs[0]


def main(caminho_pdf=None):
    if caminho_pdf is None:
        caminho_pdf = pdf_mais_recente()
    else:
        caminho_pdf = Path(caminho_pdf)

    print(f"[PDF] Lendo: {caminho_pdf.name}")

    m = re.search(r"(\d{2}-\d{2}-\d{4})", caminho_pdf.stem)
    data_str = m.group(1).replace("-", "/") if m else datetime.now().strftime("%d/%m/%Y")

    ocorrencias = extrair_ocorrencias(caminho_pdf, data_str)
    print(f"[PDF] {len(ocorrencias)} ocorrência(s) encontrada(s).")

    if not ocorrencias:
        print("[Info] Nenhuma ocorrência detectada. Nenhum arquivo gerado.")
        return None

    for oc in ocorrencias:
        print(f"  • {oc['nome']} — {oc['ocorrencia']}  |  Previsto: {oc['previsto']}  |  Registrado: {oc['registrado']}")

    nome_saida = f"justificativas_{data_str.replace('/', '-')}.docx"
    caminho_saida = PASTA_DOCX / nome_saida
    gerar_docx(ocorrencias, data_str, caminho_saida)
    return str(caminho_saida)


if __name__ == "__main__":
    main()
